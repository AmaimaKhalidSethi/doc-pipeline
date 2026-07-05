"""
Offline test suite: parsers, the recursive strict-schema builder, and the
full FastAPI request lifecycle, all without touching the network or
requiring GROQ_API_KEY. The extraction agent is mocked -- see
tests/test_sample_documents.py for the (optional, real-API) demo run
against the 5 sample documents.

Run with: pytest tests/ -v
"""
from __future__ import annotations

import io
import os
import sys
from pathlib import Path
from unittest.mock import MagicMock

import pytest

BACKEND_DIR = str(Path(__file__).parent.parent / "backend")
sys.path.insert(0, BACKEND_DIR)

os.environ.setdefault("DOC_PIPELINE_DB", ":memory:")


# --- parsers -----------------------------------------------------------

def test_extract_txt():
    from parsers import extract_text

    assert extract_text("note.txt", b"Hello world") == "Hello world"


def test_extract_docx():
    from docx import Document as DocxDocument
    from parsers import extract_text

    doc = DocxDocument()
    doc.add_paragraph("Line one")
    doc.add_paragraph("Line two")
    buf = io.BytesIO()
    doc.save(buf)
    text = extract_text("test.docx", buf.getvalue())
    assert "Line one" in text
    assert "Line two" in text


def test_extract_pdf():
    import pymupdf
    from parsers import extract_text

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from a PDF")
    pdf_bytes = doc.tobytes()
    doc.close()
    text = extract_text("test.pdf", pdf_bytes)
    assert "Hello from a PDF" in text


def test_unsupported_extension_raises():
    from parsers import UnsupportedFileType, extract_text

    with pytest.raises(UnsupportedFileType):
        extract_text("file.xyz", b"data")


# --- strict schema builder ----------------------------------------------

def test_strict_schema_flat_and_nested():
    from extraction import build_strict_schema
    from schemas import DocumentExtraction

    schema = build_strict_schema(DocumentExtraction)
    assert schema["additionalProperties"] is False
    assert set(schema["required"]) == set(schema["properties"].keys())

    for def_schema in schema.get("$defs", {}).values():
        assert def_schema["additionalProperties"] is False
        assert set(def_schema["required"]) == set(def_schema["properties"].keys())
        for prop in def_schema["properties"].values():
            assert "default" not in prop


# --- API lifecycle (extraction agent mocked) -----------------------------

@pytest.fixture()
def client(tmp_path, monkeypatch):
    db_path = tmp_path / "test.db"
    monkeypatch.setenv("DOC_PIPELINE_DB", str(db_path))

    # Reload db module so it picks up the patched env var (it reads
    # DOC_PIPELINE_DB at import time).
    import importlib
    import db as db_module
    importlib.reload(db_module)

    import main as main_module
    importlib.reload(main_module)

    from schemas import ActionItem, DocumentExtraction

    fake_extraction = DocumentExtraction(
        summary="A test summary.",
        entities=["Acme Corp", "Jane Doe"],
        key_dates=["2026-08-01"],
        key_terms=["onboarding", "contract"],
        action_items=[ActionItem(description="Sign the contract", owner="Jane Doe", due_date="2026-08-01")],
    )
    mock_agent = MagicMock()
    mock_agent.extract.return_value = fake_extraction
    main_module.get_agent = lambda: mock_agent
    main_module._agent = mock_agent

    from fastapi.testclient import TestClient

    with TestClient(main_module.app) as test_client:
        yield test_client


def test_health(client):
    resp = client.get("/health")
    assert resp.status_code == 200
    assert resp.json() == {"status": "ok"}


def test_upload_list_detail_delete_cycle(client):
    files = {"file": ("notes.txt", b"Meeting with Jane Doe on 2026-08-01.", "text/plain")}
    resp = client.post("/documents/upload", files=files)
    assert resp.status_code == 200
    body = resp.json()
    assert body["filename"] == "notes.txt"
    assert body["extraction"]["summary"] == "A test summary."
    doc_id = body["id"]

    resp = client.get("/documents")
    assert resp.status_code == 200
    assert any(d["id"] == doc_id for d in resp.json())

    resp = client.get(f"/documents/{doc_id}")
    assert resp.status_code == 200
    assert resp.json()["extraction"]["entities"] == ["Acme Corp", "Jane Doe"]

    resp = client.delete(f"/documents/{doc_id}")
    assert resp.status_code == 200

    resp = client.get(f"/documents/{doc_id}")
    assert resp.status_code == 404


def test_upload_rejects_unsupported_type(client):
    resp = client.post("/documents/upload", files={"file": ("bad.xyz", b"data", "text/plain")})
    assert resp.status_code == 400


def test_upload_rejects_empty_file(client):
    resp = client.post("/documents/upload", files={"file": ("empty.txt", b"", "text/plain")})
    assert resp.status_code == 400


def test_get_missing_document_404(client):
    resp = client.get("/documents/99999")
    assert resp.status_code == 404


def test_upload_rejects_oversized_file(client, monkeypatch):
    import main as main_module

    monkeypatch.setattr(main_module, "MAX_UPLOAD_BYTES", 10)  # tiny cap for the test
    files = {"file": ("big.txt", b"this is definitely more than ten bytes", "text/plain")}
    resp = client.post("/documents/upload", files=files)
    assert resp.status_code == 413


def test_upload_rejects_missing_filename(client):
    # Starlette's own multipart parser rejects an empty filename before our
    # handler runs at all (422); our explicit `if not file.filename` check
    # in main.py is a defensive backstop for cases where the parser lets
    # an empty/None filename through, which is version-dependent behavior.
    resp = client.post("/documents/upload", files={"file": ("", b"data", "text/plain")})
    assert resp.status_code in (400, 422)


def test_filename_is_sanitized(client):
    files = {"file": ("../../etc/passwd; rm -rf.txt", b"Meeting notes.", "text/plain")}
    resp = client.post("/documents/upload", files=files)
    assert resp.status_code == 200
    stored_name = resp.json()["filename"]
    assert ".." not in stored_name
    assert "/" not in stored_name
    assert ";" not in stored_name


def test_parse_failure_does_not_leak_exception_text(client, monkeypatch):
    def boom(*args, **kwargs):
        raise RuntimeError("some internal library path or secret detail")

    monkeypatch.setattr("main.extract_text", boom)
    resp = client.post("/documents/upload", files={"file": ("notes.txt", b"hello", "text/plain")})
    assert resp.status_code == 422
    assert "internal library path" not in resp.text
    assert "secret detail" not in resp.text


def test_upload_rate_limiting(client, monkeypatch):
    import main as main_module

    monkeypatch.setattr(main_module, "_UPLOAD_RATE_LIMIT", 2)
    main_module._upload_timestamps.clear()

    files = {"file": ("notes.txt", b"Meeting with Jane Doe on 2026-08-01.", "text/plain")}
    for _ in range(2):
        resp = client.post("/documents/upload", files=files)
        assert resp.status_code == 200

    resp = client.post("/documents/upload", files=files)
    assert resp.status_code == 429
