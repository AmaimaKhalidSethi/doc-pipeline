"""
Document parsers for PDF, DOCX, and TXT input.

Written for: pymupdf 1.27.2, python-docx 1.2.0

Research note: PyMuPDF's import name is `pymupdf`, not the legacy `fitz`
alias -- `import fitz` still works but `import pymupdf` is what current
docs recommend, so that's what's used here.
"""
from __future__ import annotations

import io

import pymupdf
from docx import Document as DocxDocument

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileType(ValueError):
    pass


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch to the right parser based on file extension and return plain text."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf_text(content)
    if lower.endswith(".docx"):
        return _extract_docx_text(content)
    if lower.endswith(".txt"):
        return _extract_txt_text(content)
    raise UnsupportedFileType(
        f"Unsupported file type for '{filename}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def _extract_pdf_text(content: bytes) -> str:
    text_parts = []
    with pymupdf.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def _extract_docx_text(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def _extract_txt_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


def file_type_from_name(filename: str) -> str:
    lower = filename.lower()
    for ext in SUPPORTED_EXTENSIONS:
        if lower.endswith(ext):
            return ext.lstrip(".")
    raise UnsupportedFileType(filename)
